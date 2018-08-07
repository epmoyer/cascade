# Linter Directives
#     General
#         pylint: disable=locally-disabled, line-too-long, import-error, no-name-in-module
#         pylint: disable=locally-disabled, too-many-locals, too-many-branches, too-many-statements
#         pylint: disable=locally-disabled, too-many-instance-attributes

# Standard library
import os
import sys
import json
import logging
import copy
from collections import OrderedDict
import pytest

# Libraries
from docx import Document

# Local
from cascade.quicklog import Quicklog

# Get the path to the test directory (this file's path)
test_root_path = os.path.dirname(os.path.realpath(__file__))

# Start default logger (imports below will use it)
qlog = Quicklog(
    log_filename=os.path.join(
        test_root_path,
        'results',
        'log.txt'),
    logging_level=logging.DEBUG)
qlog.begin(os.path.dirname(__file__))
qlog.info('test_check.py')

from cascade.word_docx import WordDocx
from cascade.cmd_check import check
from cascade.custom_exceptions import FatalUserError

DUMMY_TEXT = (
    'Lorem ipsum dolor sit amet, consectetur adipiscing elit. Praesent iaculis' +
    ' ante tortor, eget vestibulum eros elementum vel. Etiam sollicitudin magna' +
    ' placerat, placerat elit non, tempus risus. Vivamus sed ullamcorper eros.')


@pytest.mark.parametrize("doc_type, expected_result", [
    # pylint: disable=locally-disabled, bad-whitespace

    # Positives
    ('normal',                           True),
    ('normal_unassigned',                True),

    # Negatives
    ('no_document_info_directive',       False),
    ('missing_object_id_number',         False),
    ('repeat_object_id_number',          False),
    ('object_id_number_exceeds_next_id', False),
    ('bad_object_id_number',             False),
    ('bad_object_id_prefix',             False),
    ('bad_directive_json',               False),
    ('unexpected_json_field',            False),
    ('missing_prefix',                   False),
    ('missing_next_id',                  False),
    ('old_object_id_format',             False),
])

def test_check(doc_type, expected_result):
    test_doc_filename = build_test_document(doc_type)
    arguments = {'<requirements.docx>': os.path.join(test_root_path, 'results', test_doc_filename)}
    result = False

    if doc_type == 'bad_directive_json':
        with pytest.raises(FatalUserError):
            result = check(arguments)
    else:
        result = check(arguments)

    assert result == expected_result

def build_test_document(doc_type):
    with open(os.path.join(test_root_path, 'document_info.json')) as data_file:    
        document_info_dict = json.load(data_file)
    if doc_type == 'missing_prefix':
        del document_info_dict['#document_info']['object_ids'][0]['prefix']
    if doc_type == 'missing_next_id':
        del document_info_dict['#document_info']['object_ids'][0]['next_id']

    # Create base document
    doc_filename = os.path.join(test_root_path, 'results', 'test_base.docx')
    document = Document(os.path.join(test_root_path,'hidden_style_master.docx'))
    document.add_heading('TEST: check')
    document.add_paragraph('TEST CASE: {}'.format(doc_type))
    document.add_paragraph('First content paragraph.')
    document.add_paragraph('Last content paragraph.')
    document.save(doc_filename)

    # Manipulate base document using WordDocx class
    doc = WordDocx(qlog, doc_filename)
    p = doc.paragraphs[-2] # Get next to last paragraph
    requirement_directive_dict = OrderedDict([
        ('id', 'ABC-DEF-000'),
        ('method', 'I'),
        ('old_id', 'SYS-795')
        ])

    if doc_type != 'no_document_info_directive':
        p = doc.insert_directive(p, document_info_dict)
    p = doc.insert_paragraph(p, DUMMY_TEXT)

    for i in range(3):
        publish_dict = copy.copy(requirement_directive_dict)
        id_prefix = '-'.join(publish_dict['id'].split('-')[:-1]) + '-'
        if i == 1:
            # Inject requirement directive errors
            if doc_type == 'normal_unassigned':
                publish_dict['id'] = id_prefix + '?'
            if doc_type == 'missing_object_id_number':
                publish_dict['id'] = id_prefix
            if doc_type == 'bad_object_id_number':
                publish_dict['id'] = id_prefix + 'X'
            if doc_type == 'bad_object_id_prefix':
                publish_dict['id'] = 'ABC-QQQ-' + publish_dict['id'].split('-')[-1]
            if doc_type == 'object_id_number_exceeds_next_id':
                publish_dict['id'] = id_prefix + '101'
            if doc_type == 'unexpected_json_field':
                publish_dict['unexpected'] = '1234'

            if doc_type == 'bad_directive_json':
                p = doc.insert_paragraph(p, '${"id":"ABC-DEF-001", "metho}$')
            elif doc_type == 'old_object_id_format':
                p = doc.insert_paragraph(p, '[ABC-DEF-123, X]')
            else:
                p = doc.insert_directive(p, publish_dict, simple=True, format_type='directive_visible')
        else:
            p = doc.insert_directive(p, publish_dict, simple=True, format_type='directive_visible')

        if not doc_type == 'repeat_object_id_number':
            incerment_requirement_directive_dict(requirement_directive_dict)
        p = doc.insert_paragraph(p, DUMMY_TEXT)

    output_filename = 'test_check_{}.docx'.format(doc_type)
    doc.save(os.path.join(test_root_path, 'results', output_filename))
    return output_filename

def incerment_requirement_directive_dict(requirement_directive_dict):
    object_id = requirement_directive_dict['id']
    object_id_pieces = object_id.split('-')
    object_id_pieces[-1] = '{:03d}'.format(int(object_id_pieces[-1]) + 1)
    object_id = '-'.join(object_id_pieces)
    requirement_directive_dict['id'] = object_id
