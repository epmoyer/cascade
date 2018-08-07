"""Library for manipulating Microsoft Word files using python-docx

The python-docx library is documented at:
   https://python-docx.readthedocs.io/en/latest/
"""

# Standard library
from collections import namedtuple, defaultdict
import pprint

# Libraries
from docx import Document
from docx.shared import Pt as Pt
from docx.shared import RGBColor as RGBColor
# The following imports specifically support the iter_block_items() function
# as documented in https://github.com/python-openxml/python-docx/issues/40
from docx.document import Document as docx_Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

# Local
from cascade.util import (
    make_json, json_to_dict, is_shortform_dict, make_json_autoformat,
    extract_json_from_directive, expand_shortform_dict, get_directive_type,
    to_snippet)
from cascade.custom_exceptions import FatalUserError

pp = pprint.PrettyPrinter(indent=3)

class WordDocx(object):
    def __init__(self, qlog, filename):
        self._qlog = qlog
        self._filename_original = filename
        self._document = Document(filename)
        self.paragraphs = self._document.paragraphs
        self.format_types = ['directive_visible', 'directive_hidden']
        self._Directive = namedtuple('Directive', 'paragraphs as_dict')
        self._directives = []
        self.requirements = []
        self.doc_info_directive = None
        self.find_directives()

    def resync_paragraphs(self):
        ''' Get local copy of paragraphs again

        This needs to be done if paragraphs have changed and we need to 
        walk the list of paragraphs again.  A local copy of the 
        paragraphs list needs to be maintained (in general), otherwise
        comparisons of paragraph objects will fail when iterating over
        the list of paragraphs.
        '''
        self.paragraphs = self._document.paragraphs

    def find_directives(self):
        '''Find all directives in document'''
        self._qlog.debug("Finding directives...")
        directive_in_progress = False
        directive_paragraphs = []
        current_requirement = {}

        json_text = ""
        current_heading_text = ""
        for p in self.paragraphs:
            p_added = False
            if self.get_heading_level(p):
                p_is_heading = True
                current_heading_text = p.text.strip()
            else:
                p_is_heading = False

            if not directive_in_progress:
                if '${' in p.text:
                    # Start found
                    directive_in_progress = True
                    directive_paragraphs.append(p)
                    json_text += p.text
                    p_added = True

                if current_requirement:
                    # A feature is in progress.  Add to it or commit it.
                    if p_is_heading:
                        # Current paragraph is a heading 
                        # so commit the feature in progress
                        self.requirements.append(current_requirement)
                        current_requirement = {}
                    else:
                        if p.text and not '${' in p.text:
                            # Current paragraph is requirements text, so add it to the feature
                            # in progress
                            if current_requirement['req_text']:
                                current_requirement['req_text'] += '\n' + p.text
                            else:
                                current_requirement['req_text'] += p.text
            if directive_in_progress:
                if not p_added:
                    # Capture constituent paragraphs
                    directive_paragraphs.append(p)
                    json_text += p.text
                    p_added = True

                if '}$' in p.text:
                    # End found
                    json_text = json_text.replace('“', '"').replace('”', '"').strip('$')
                    self._qlog.debug('Directive JSON: "{}"'.format(json_text))
                    as_dict = json_to_dict(json_text)
                    if as_dict:
                        directive = self._Directive(
                            directive_paragraphs,
                            as_dict
                            )
                        self._directives.append(directive)
                        if '#document_info' in as_dict:
                            self.doc_info_directive = directive
                        self._qlog.debug("Found directive {}".format(as_dict))
                    else:
                        raise FatalUserError("JSON error")

                    # If this is a requirements directive
                    if is_shortform_dict(as_dict):
                        # If the previous requirement has not been committed, then
                        # commit it now.
                        if current_requirement:
                            self.requirements.append(current_requirement)
                            current_requirement = {}

                        # Create requirement dict
                        current_requirement = {
                            'directive': directive,
                            'req_text': '',
                            'heading_text': current_heading_text
                        }

                    # Clear search
                    directive_in_progress = False
                    directive_paragraphs = []
                    json_text = ''

        # Commit the last feature
        if current_requirement:
            self.requirements.append(current_requirement)
            current_requirement = {}

        self._qlog.debug("Found {} directives".format(len(self._directives)))
        self._qlog.debug("Found {} requirements".format(len(self.requirements)))
        self._qlog.debug("Requirements:\n {}".format(pp.pformat(self.requirements)))

    def get_text(self, paragraph):
        '''Get a paragraph's text'''
        #pylint: disable=locally-disabled, no-self-use

        return paragraph.text

    def rewrite_directive(self, directive):
        '''Rewrite a (presumably modified) directive to the document'''
        # Insert new directive paragraph(s)
        first_p = directive.paragraphs[0]
        directive_json = make_json_autoformat(directive.as_dict)
        if is_shortform_dict(directive.as_dict):
            paragraph = first_p.insert_paragraph_before(directive_json)
            self.format_paragraph(paragraph, 'directive_visible')
        else:
            lines = directive_json.split('\n')
            for line in lines:
                paragraph = first_p.insert_paragraph_before(line)
                self.format_paragraph(paragraph, 'directive_hidden')
        # Delete old directive paragraph(s)
        for paragraph in directive.paragraphs:
            self.delete_paragraph(paragraph)

    def delete_paragraph(self, paragraph):
        '''See https://github.com/python-openxml/python-docx/issues/33
        From 3/16/2015 post by docx author (scanny)
        Caveat: Paragraph must be "simple" (i.e not contain tables or images),
        otherwise that content will become "zombie" content consuming doc space
        but being otherwise un-viewable / un-editable / un-deletable.
        '''
        #pylint: disable=locally-disabled, no-self-use, protected-access, invalid-name

        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
        # In spite of scanny's post, I strongly suspect the correct line of code
        # is the one below, rather than the one above.
        #paragraph._p = paragraph._element = None

    def delete_table_row(self, table, row):
        '''See https://github.com/python-openxml/python-docx/issues/83
        '''
        #pylint: disable=locally-disabled, no-self-use, protected-access, invalid-name
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)

    def iter_block_items(self, parent):
        '''See https://github.com/python-openxml/python-docx/issues/40
        Yield each paragraph and table child within *parent*, in document order.
        Each returned value is an instance of either Table or Paragraph. *parent*
        would most commonly be a reference to a main Document object, but
        also works for a _Cell object, which itself can contain paragraphs and tables.

        Commentary: 
           Cascade uses this function to walk through the Paragraphs and Tables
           of a document in order.  It is (currently) the only way in the 
           python-docx API to determine the physical location of tables
           within a document.
        '''
        if isinstance(parent, docx_Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("Something is not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def format_paragraph(self, paragraph, format_type):
        '''Apply formatting to a paragraph'''
        if format_type not in self.format_types:
            raise ValueError('format_type "{}" did not match one of {}'.format(format_type, self.format_types))
        if format_type == 'directive_visible':
            paragraph.style = self._document.styles['Cascade Directive']
        elif format_type == 'directive_hidden':
            paragraph.style = self._document.styles['Cascade Hidden Directive']

    def get_heading_level(self, paragraph):
        '''Numerical heading level of paragraph if is heading.  None otherwise'''
        #pylint: disable=locally-disabled, no-self-use

        style = paragraph.style.name
        if style.startswith('Heading'):
            # Style has form 'Heading N'
            return int(style.split(' ')[1])
        elif style.startswith('Appendix_') and '_Level_' in style:
            # Style has form 'Appendix_X_Level_N'
            return int(style.split('_')[-1])
        return None

    def insert_paragraph(self, paragraph, text):
        '''Insert a new paragraph AFTER the specified paragraph'''
        next_pararaph, index = self.get_next_paragrah(paragraph)
        if next_pararaph:
            new_paragraph = next_pararaph.insert_paragraph_before(text)
            self.paragraphs.insert(index, new_paragraph)
            return new_paragraph
        else:
            self._qlog.critical('WordDocx.insert_paragraph() cannot be called on last paragraph.')


    def insert_directive(self, paragraph, data_dict, simple=False, format_type='directive_hidden'):
        '''Insert dict as JSON directive located AFTER the specified paragraph

        The directive will be formatted using the specified format_type
        The directive will be indented to match the specified paragraph.

        Returns the new paragraph
        '''
        insert_text = make_json(data_dict, simple)
        self._qlog.debug('Inserting directive: "{}"'.format(insert_text))
        next_pararaph, index = self.get_next_paragrah(paragraph)
        if next_pararaph:
            new_paragraph = next_pararaph.insert_paragraph_before(insert_text)
            self.format_paragraph(new_paragraph, format_type)
            new_paragraph.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
            self.paragraphs.insert(index, new_paragraph)
            return new_paragraph
        else:
            self._qlog.error('The following directive could not be inserted because it would have '+
                             'been the last paragraph in the document, and directives should never be located ' +
                             'at the very end of a document. Something is wrong. : "{}"'.format(insert_text))

    def get_next_paragrah(self, paragraph):
        '''Given a paragraph, return the next paragraph in the doc & the index of the paragraph

        Returns a tuple: (paragraph, index)
            paragraph and index will be null if no solution was found
        python-docx has no paragraph.next() method.  This method implements a brute-force search
        which is adequate at present.
        '''
        found = False
        #self._qlog.debug('get_next_paragraph: Searching for: "{}"'.format(to_snippet(paragraph.text)))
        for index, search_p in enumerate(self.paragraphs):
            #self._qlog.debug('get_next_paragraph: Comparing: "{}"'.format(to_snippet(search_p.text)))
            if found:
                return (search_p, index)
            if search_p == paragraph:
                found = True
        if not found:
            self._qlog.debug('get_next_paragraph: Search target not found.')
        else:
            self._qlog.debug('get_next_paragraph: Search target was last paragraph in doc.')
        return (None, None)

    def get_clusters(self):
        '''Get a list containing paragraph clusters.

        A cluster is a group of paragraphs (1 or more) which make up a:
            heading, directive, or body

        Each list element is a dict with one of the following forms:
            HEADING CLUSTER
                {
                    'cluster_type':   'heading',
                    'paragraphs':     <a list containing the heading paragraph>,
                    'heading_level':  <an integer>
                }

            BODY CLUSTER
                {
                    'cluster_type':   'body',
                    'paragraphs':     <a list containing the body paragraphs (1 or more)>,
                }

            DIRECTIVE CLUSTER
                {
                    'cluster_type':   'directive',
                    'paragraphs':     <a list containing the directive paragraphs (1 or more)>,
                    'directive':      <a dict representing the directive data>
                }
        '''
        clusters = []
        directive_in_progress = False
        cluster_paragraphs = []
        json_text = ''
        for paragraph in self.paragraphs:
            p_added = False
            if not directive_in_progress:
                if '${' in paragraph.text:
                    # Start found
                    directive_in_progress = True
                    cluster_paragraphs.append(paragraph)
                    json_text += paragraph.text
                    p_added = True
                else:
                    #-- if paragraph.heading_level:
                    heading_level = self.get_heading_level(paragraph)
                    if heading_level:
                            clusters.append({
                                'cluster_type':   'heading',
                                'paragraphs':    [paragraph],
                                'heading_level': heading_level 
                            })
                    else:
                        clusters.append({
                                'cluster_type': 'body',
                                'paragraphs': [paragraph]
                            })
            if directive_in_progress:
                if not p_added:
                    # Capture constituent paragraphs
                    cluster_paragraphs.append(paragraph)
                    json_text += paragraph.text
                    p_added = True
                if '}$' in paragraph.text:
                    # End of directive found
                    json_text = extract_json_from_directive(json_text)
                    self._qlog.debug('Directive JSON: "{}"'.format(json_text))
                    as_dict = json_to_dict(json_text)
                    if as_dict:
                        if is_shortform_dict(as_dict):
                            as_dict = expand_shortform_dict(as_dict)
                        clusters.append({
                            'cluster_type': 'directive',
                            'paragraphs':   cluster_paragraphs,
                            'directive':    as_dict
                        })
                        self._qlog.debug("Found directive {}".format(as_dict))
                    else:
                        raise RuntimeError("JSON error")

                    # Clear search
                    directive_in_progress = False
                    cluster_paragraphs = []
                    json_text = ''

        if self._qlog.debug_is_enabled():
            # Log the fist handful of clusters for debug
            max_log_clusters = 200
            self._qlog.debug('Logging {} of {} clusters:'.format(min(len(clusters), max_log_clusters), len(clusters)))
            self._qlog.debug('clusters:\n' + pp.pformat(clusters[:max_log_clusters]))

            # Log the cluster stats
            stat_cluster_types = defaultdict(int)
            stat_directive_types = defaultdict(int)
            for cluster in clusters:
                stat_cluster_types[cluster['cluster_type']] += 1
                if cluster['cluster_type'] == 'directive':
                    directive_type = get_directive_type(cluster['directive'])
                    stat_directive_types[directive_type] += 1
            self._qlog.debug('Cluster stats:\nCluster Types:{}\nCluster directive types:{}'.format(
                stat_cluster_types,
                stat_directive_types))

        return clusters

    def save(self, filename):
        '''Save the current file'''
        # TODO: Should this close as well?
        self._document.save(filename)
