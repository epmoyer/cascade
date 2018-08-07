"""Handler for the 'debug-dump' command
(Commands are issued on the command line, per the docopt syntax in __main__.py)
"""

# Standard library
import os
import sys

# Libraries
from docx.table import Table
from docx.text.paragraph import Paragraph

# Local
from cascade.word_docx import WordDocx
from cascade.util import Indent, get_heading_level
from cascade import quicklog
from cascade.util_eliot import log_function

qlog = quicklog.get_logger()
lprint = qlog.lprint

# pylint: disable=line-too-long, protected-access

def shorten(text):
    display_chars_max = 60
    short_text = text if len(text) < display_chars_max else text[:display_chars_max] + '...'
    # TODO: Handle Unicode to stdout.  For now, unprintable Unicode is replaced with "?" using .encode() with errors='replace'
    # See: http://stackoverflow.com/questions/14630288/unicodeencodeerror-charmap-codec-cant-encode-character-maps-to-undefined
    # print("stdout encoding:{}".format(sys.stdout.encoding))
    return short_text.encode(sys.stdout.encoding, errors='replace').decode(sys.stdout.encoding)

@log_function
def dump(arguments):
    '''Dump the details of a word .docx file, as interpreted by python-docx via the WordDocx module
    
    (For development debugging)
    '''

    filename = arguments['<requirements.docx>']
    if not os.path.isfile(filename):
        qlog.error('The file "{}" does not exist'.format(filename))
        return

    lprint('Dumping document "{}"...'.format(filename))
    doc = WordDocx(qlog, os.path.abspath(filename))

    #-------------------------
    # iter_block_items()
    #-------------------------
    
    print('BLOCK ITEMS:')
    previous = None
    for item in doc.iter_block_items(doc._document):
        if isinstance(item, Table):
            if isinstance(previous, Paragraph):
                # Show paragraph containing the table
                print('   Paragraph: "{}"'.format(previous.text))
            print('     Table.')
        previous = item

    #-------------------------
    # Paragraphs
    #-------------------------

    print('PARAGRAPHS:')
    indent = Indent()
    indent.inc()
    indent.print("Num | Heading    |Level | Indent | Text")
    indent.print("----|------------|------|--------|---------------------  ")

    for p_number, paragraph in enumerate(doc.paragraphs):
        style = paragraph.style.name
        p_indent = paragraph.paragraph_format.left_indent
        indent_inches = 0 if p_indent is None else p_indent.inches
        heading_level = get_heading_level(style)
        indent.print("{:04d}| {:>10} | {:>4} | {: 0.3f} | {}".format(p_number, style, str(heading_level), indent_inches, shorten(paragraph.text)))
    indent.dec()

    #-------------------------
    # Tables
    #-------------------------

    print('\nTABLES:')
    indent = Indent()
    indent.inc()
    # Note: This code accesses the private ._document property and walks tables with
    #       intimate knowledge of the python-docx table structure.  Tables (as of
    #       this writing) are not abstracted by the WordDocx module, so we will
    #       access them using brute force.
    indent.print('Found {} tables...'.format(len(doc._document.tables)))
    for table_index, table in enumerate(doc._document.tables):
        indent.print('TABLE {:03d}:'.format(table_index))
        rows = table.rows
        indent.inc()
        indent.print('Found {} rows...'.format(len(rows)))
        for row_index, row in enumerate(rows):
            indent.print('ROW {:03d}:'.format(row_index))
            indent.inc()
            for cell_index, cell in enumerate(row.cells):
                indent.print('CELL {:03d}:'.format(cell_index))
                indent.inc()
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    indent.print("PARAGRAPH {:03d}: {}".format(paragraph_index, shorten(paragraph.text)))
                indent.dec()
            indent.dec()
        indent.dec()
    indent.dec()

    print("\nDone.")
